import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import os
import re
from io import BytesIO

# --- LIBRERIE STAMPA ---
try:
    from docx import Document
    from docx.shared import Inches
    import plotly.io as pio
    DOC_OK = True
except:
    DOC_OK = False

# --- LOGICA DI CALCOLO ORIGINALE (VBA STYLE) ---
@st.cache_data(show_spinner=False)
def elaborazione_vba_originale(df_values, l_barra, n_sigma, limit_val):
    # Calcolo mm: L * sin(rad(deg))
    data_mm = l_barra * np.sin(np.radians(df_values.values))
    # Delta C0 (Rispetto alla riga 0)
    data_c0 = data_mm - data_mm[0, :]
    # Deformata Cumulata (Somma vettoriale dei sensori in riga)
    data_cp0 = np.cumsum(data_c0, axis=1)
    
    # Filtro Limiti Hard
    data_cp0[np.abs(data_cp0) > limit_val] = np.nan
    
    # Filtro Gauss (Sigma)
    means = np.nanmean(data_cp0, axis=0)
    stds = np.nanstd(data_cp0, axis=0)
    for j in range(data_cp0.shape[1]):
        m, s = means[j], stds[j]
        mask = (data_cp0[:, j] < m - n_sigma*s) | (data_cp0[:, j] > m + n_sigma*s) | (np.isnan(data_cp0[:, j]))
        data_cp0[mask, j] = m
        
    return pd.DataFrame(data_cp0, index=df_values.index).ffill().fillna(0)

# --- FUNZIONE RICHIAMATA DA APP_DIMOS.PY ---
def run_elettrolivelle_advanced():
    st.header("📏 Modulo Elettrolivelle")

    with st.sidebar:
        st.divider()
        st.subheader("⚙️ Setup Dati")
        file_input = st.file_uploader("Carica Excel Monitoraggio", type=['xlsm', 'xlsx'], key="el_up")
        
        if file_input:
            asse_sel = st.selectbox("Asse", ["X", "Y", "Z"], key="el_asse")
            l_barra = st.number_input("Lunghezza Barra (mm)", value=3000)
            sigma_val = st.slider("Filtro Sigma", 1.0, 4.0, 2.0)
            limit_val = st.number_input("Soglia Errore (mm)", value=30.0)
            
            st.divider()
            st.subheader("🎬 Opzioni Video")
            step_video = st.select_slider("Campionamento:", 
                options=["Ogni Lettura", "1 Giorno", "1 Settimana"], value="1 Giorno")
            vel_animazione = st.slider("Velocità (ms)", 100, 1000, 400)

    if not file_input:
        st.info("Carica il file Excel per procedere.")
        return

    # Lettura Excel
    xls = pd.ExcelFile(file_input)
    # Filtriamo i fogli dati (ETS_...)
    sheets = [s for s in xls.sheet_names if s.startswith("ETS_") and not s.endswith(("C0", "CP0"))]
    
    tab1, tab2 = st.tabs(["📊 Grafico Animato", "🖨️ Esportazione Report"])

    with tab1:
        sel_sheet = st.selectbox("Seleziona Stringa", sheets)
        df_full = pd.read_excel(file_input, sheet_name=sel_sheet)
        df_full.columns = [str(c).strip() for c in df_full.columns]
        time_col = pd.to_datetime(df_full.iloc[:, 0])

        # --- LOGICA ARRAY: SEQUENZA FISICA ---
        sensor_order = []
        if "ARRAY" in xls.sheet_names:
            df_arr = pd.read_excel(file_input, sheet_name="ARRAY", header=None)
            riga = df_arr[df_arr[0] == sel_sheet]
            if not riga.empty:
                sensor_order = riga.iloc[0, 1:].dropna().astype(str).tolist()

        # Selezione colonne sensori in ordine ARRAY
        cols_found = []
        labels_x = []
        for s_id in sensor_order:
            pattern = rf"{s_id}.*_{asse_sel}"
            match = [c for c in df_full.columns if re.search(pattern, str(c), re.IGNORECASE)]
            if match:
                cols_found.append(match[0])
                labels_x.append(s_id)

        if not cols_found:
            st.error(f"Nessun dato per asse {asse_sel} nel foglio {sel_sheet}")
            return

        # Calcolo Deformata Cumulata
        df_cp0 = elaborazione_vba_originale(df_full[cols_found].ffill(), l_barra, sigma_val, limit_val)
        
        # Campionamento per fluidità video
        df_calc = df_cp0.copy()
        df_calc['Data_Ora'] = time_col
        if step_video == "1 Giorno":
            df_sampled = df_calc.groupby(df_calc['Data_Ora'].dt.date).first().drop(columns='Data_Ora')
        elif step_video == "1 Settimana":
            df_sampled = df_calc.set_index('Data_Ora').resample('W').first().dropna()
        else:
            df_sampled = df_calc.set_index('Data_Ora')

        # --- GRAFICO ANIMATO ---
        fig = go.Figure()
        fig.add_trace(go.Scatter(
            x=labels_x, y=df_sampled.iloc[0],
            mode='lines+markers+text',
            text=[f"{v:.2f}" for v in df_sampled.iloc[0]],
            textposition="top center",
            line=dict(color="#1f77b4", width=2),
            marker=dict(size=12, color="blue", line=dict(width=1, color="black"))
        ))

        fig.update_layout(
            xaxis=dict(type='category', title="Sequenza Fisica"),
            yaxis=dict(range=[-limit_val-5, limit_val+5], title="Spostamento (mm)"),
            template="plotly_white", height=600,
            sliders=[{"active": 0, "steps": [{"method": "animate", "label": str(t), 
                       "args": [[str(i)], {"frame": {"duration": vel_animazione, "redraw": True}}]} 
                      for i, t in enumerate(df_sampled.index)]}]
        )

        fig.frames = [go.Frame(data=[go.Scatter(
            x=labels_x, y=df_sampled.iloc[i],
            text=[f"{v:.2f}" for v in df_sampled.iloc[i]],
            marker=dict(color=['red' if abs(v)>10 else 'green' for v in df_sampled.iloc[i]])
        )], name=str(i)) for i in range(len(df_sampled))]

        st.plotly_chart(fig, use_container_width=True)

    with tab2:
        st.subheader("Esportazione Documento Word")
        if st.button("🚀 GENERA DOCUMENTAZIONE"):
            if not DOC_OK:
                st.error("Libreria docx mancante.")
            else:
                with st.spinner("Generazione..."):
                    doc = Document()
                    doc.add_heading(f"Report Deformata - {sel_sheet}", 0)
                    
                    # Immagine del frame corrente
                    buf = BytesIO()
                    pio.write_image(fig, buf, format="png", width=800, height=450)
                    doc.add_picture(buf, width=Inches(6))
                    
                    target = BytesIO()
                    doc.save(target)
                    st.download_button("📥 Scarica File", target.getvalue(), f"Report_{sel_sheet}.docx")
