import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import re
import os
from io import BytesIO

# --- LIBRERIE STAMPA ---
try:
    from docx import Document
    from docx.shared import Inches
    import plotly.io as pio
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def run_plotter():
    # Logo DIMOS
    col_l, _ = st.columns([1, 2])
    with col_l:
        if os.path.exists("logo_dimos.jpg"): st.image("logo_dimos.jpg", width=300)
    
    st.header("📈 Plotter DIMOS - Analisi e Stampe")

    file_input = st.file_uploader("Carica Excel Monitoraggio", type=['xlsx', 'xlsm'], key="p_up")
    if not file_input: return

    xls = pd.ExcelFile(file_input)
    
    # --- 1. MAPPATURA GRANULARE (NAME LAYER) ---
    anagrafica = {} 
    if "NAME" in xls.sheet_names:
        df_name = pd.read_excel(xls, sheet_name="NAME", header=None)
        # Saltiamo la prima colonna (etichette riga) e iteriamo
        for c_idx in range(1, df_name.shape[1]):
            dl = str(df_name.iloc[0, c_idx]).strip()
            sens = str(df_name.iloc[1, c_idx]).strip()
            full_web = str(df_name.iloc[2, c_idx]).strip()
            
            if dl == "nan" or full_web == "nan": continue
            if dl not in anagrafica: anagrafica[dl] = {}
            if sens not in anagrafica[dl]: anagrafica[dl][sens] = {}
            
            # Identifichiamo la grandezza (es. X [°], LM [m]) usando la tua logica split
            grandezza = full_web.split("_")[-1] if "_" in full_web else full_web
            anagrafica[dl][sens][grandezza] = full_web
    else:
        st.error("Foglio NAME non trovato. Verificare il file."); return

    # --- 2. CARICAMENTO DATI ---
    sheets_dati = [s for s in xls.sheet_names if s not in ["NAME", "ARRAY", "Info"]]
    sel_sheet = st.selectbox("Seleziona Layer Dati", sheets_dati)
    df = pd.read_excel(xls, sheet_name=sel_sheet)
    df.columns = [str(c).strip() for c in df.columns]
    if 'Data e Ora' in df.columns:
        df['Data e Ora'] = pd.to_datetime(df['Data e Ora'])
    else:
        st.error("Colonna 'Data e Ora' non trovata!"); return

    # --- 3. SIDEBAR: FILTRI E TREND ---
    st.sidebar.header("🛠️ Strumenti di Analisi")
    clean_zeros = st.sidebar.checkbox("Rimuovi Zeri (0.00)", value=True)
    use_gauss = st.sidebar.checkbox("Filtro Gaussiano (Outliers)", value=True)
    sigma_val = st.sidebar.slider("Sigma Gauss", 1.0, 5.0, 2.0, 0.1)
    
    st.sidebar.divider()
    st.sidebar.subheader("📈 Curva di Tendenza")
    grado_poly = st.sidebar.selectbox("Grado Polinomiale", [0, 2, 3, 4], 
                                     format_func=lambda x: "OFF" if x==0 else f"Grado {x}")

    # --- 4. SELEZIONE MULTIPLA (3 LIVELLI) ---
    st.divider()
    t_min, t_max = df['Data e Ora'].min().to_pydatetime(), df['Data e Ora'].max().to_pydatetime()
    sel_range = st.slider("Seleziona Periodo", t_min, t_max, (t_min, t_max))
    df_f = df[(df['Data e Ora'] >= sel_range[0]) & (df['Data e Ora'] <= sel_range[1])].copy()

    c1, c2, c3 = st.columns(3)
    with c1: 
        sel_dls = st.multiselect("1. Datalogger", sorted(anagrafica.keys()))
    
    opts_s = [f"{d} | {s}" for d in sel_dls for s in anagrafica[d].keys()]
    with c2: 
        sel_sens_pairs = st.multiselect("2. Sensori", opts_s)
    
    opts_g = set()
    for it in sel_sens_pairs:
        d, s = it.split(" | ")
        opts_g.update(anagrafica[d][s].keys())
    with c3: 
        # Qui ora vedrai X [°], Y [°], Z [°], T1 [°C], ecc. per SD_301
        sel_grands = st.multiselect("3. Grandezze Fisiche", sorted(list(opts_g)))

    # --- 5. ELABORAZIONE CON DIAGNOSTICA ---
    def process_data(col):
        y_raw = df_f[col].copy()
        z_removed = (y_raw == 0).sum()
        
        if clean_zeros: y_raw = y_raw.replace(0, np.nan)
        
        g_removed = 0
        if use_gauss and y_raw.notna().sum() > 5:
            m, s = y_raw.mean(), y_raw.std()
            mask = abs(y_raw - m) > (sigma_val * s)
            g_removed = mask.sum()
            y_raw = y_raw.mask(mask, np.nan)
        
        y_clean = y_raw.interpolate(limit_direction='both')
        
        y_trend = None
        if grado_poly > 0 and y_clean.notna().any():
            x_idx = np.arange(len(y_clean))
            valid = y_clean.notna()
            coeffs = np.polyfit(x_idx[valid], y_clean[valid], grado_poly)
            p = np.poly1d(coeffs)
            y_trend = pd.Series(p(x_idx), index=y_clean.index)
            
        return y_clean, y_trend, z_removed, g_removed

    # --- 6. VISUALIZZAZIONE E REPORT ---
    fig = go.Figure()
    diag_rows = []

    if sel_sens_pairs and sel_grands:
        for it in sel_sens_pairs:
            d, s = it.split(" | ")
            for g in sel_grands:
                if g in anagrafica[d][s]:
                    c_real = anagrafica[d][s][g]
                    if c_real in df.columns:
                        y_val, y_tr, z_rem, g_rem = process_data(c_real)
                        
                        # Aggiunta Grafico
                        fig.add_trace(go.Scatter(x=df_f['Data e Ora'], y=y_val, name=f"{s} - {g}"))
                        if y_tr is not None:
                            fig.add_trace(go.Scatter(x=df_f['Data e Ora'], y=y_tr, 
                                                   name=f"Trend {s} ({g})", line=dict(dash='dash')))
                        
                        # Dati Diagnostica
                        diag_rows.append({
                            "Sensore": s, "Grandezza": g, 
                            "Zeri Rimossi": z_rem, "Outlier (Gauss)": g_rem
                        })

        fig.update_layout(template="plotly_white", height=600, hovermode="x unified", legend=dict(orientation="h", y=-0.2))
        st.plotly_chart(fig, use_container_width=True)

        # Tabella Diagnostica a video
        st.subheader("📋 Diagnostica Qualità Dati")
        df_diag = pd.DataFrame(diag_rows)
        st.table(df_diag)

        # --- 7. STAMPE SPECULARI ---
        st.divider()
        st.subheader("💾 Esportazione Speculare")
        cw, ce = st.columns(2)

        with cw:
            if st.button("📝 GENERA REPORT WORD"):
                doc = Document()
                doc.add_heading('REPORT MONITORAGGIO DIMOS', 0)
                # Immagine
                img_io = BytesIO(pio.to_image(fig, format="png", width=1000, height=500))
                doc.add_picture(img_io, width=Inches(6.2))
                # Tabella Diagnostica nel Word
                doc.add_heading('Dettaglio Pulizia Dati', level=2)
                table = doc.add_table(rows=1, cols=3); table.style = 'Table Grid'
                hdr = table.rows[0].cells
                hdr[0].text = 'Sensore/Grandezza'; hdr[1].text = 'Zeri'; hdr[2].text = 'Gauss'
                for r in diag_rows:
                    row = table.add_row().cells
                    row[0].text = f"{r['Sensore']} - {r['Grandezza']}"
                    row[1].text = str(r['Zeri Rimossi']); row[2].text = str(r['Outlier (Gauss)'])
                
                buf_w = BytesIO(); doc.save(buf_w)
                st.download_button("📥 Scarica Word", buf_w.getvalue(), "Report_DIMOS.docx")

        with ce:
            if st.button("📊 GENERA EXCEL DATI"):
                df_out = pd.DataFrame({'Data Ora': df_f['Data e Ora']})
                for it in sel_sens_pairs:
                    d, s = it.split(" | ")
                    for g in sel_grands:
                        if g in anagrafica[d][s]:
                            y_v, y_t, _, _ = process_data(anagrafica[d][s][g])
                            df_out[f"{s}_{g}"] = y_v
                            if y_t is not None: df_out[f"Trend_{s}_{g}"] = y_t
                buf_e = BytesIO()
                df_out.to_excel(buf_e, index=False)
                st.download_button("📥 Scarica Excel", buf_e.getvalue(), "Dati_Monitoraggio.xlsx")

if __name__ == "__main__":
    run_plotter()
