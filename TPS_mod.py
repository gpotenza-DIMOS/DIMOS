# =========================================================
# DIMOS - ANALISI TOPOGRAFICA AVANZATA
# VERSIONE PROFESSIONALE MULTILAYER
# STREAMLIT CLOUD READY
# =========================================================

import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import warnings
import logging
import tempfile
import os

from docx import Document
from docx.shared import Inches

# =========================================================
# WARNING
# =========================================================
warnings.filterwarnings("ignore")

# =========================================================
# LOGGER
# =========================================================
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)

logger = logging.getLogger("DIMOS")

# =========================================================
# CONFIG PAGINA
# =========================================================
try:
    st.set_page_config(
        page_title="DIMOS - Analisi Topografica",
        layout="wide"
    )
except:
    pass

# =========================================================
# FUNZIONE CONVERSIONE NUMERICA
# =========================================================
def converti_numerico(serie):

    return pd.to_numeric(
        serie.astype(str)
        .str.replace(",", ".", regex=False)
        .str.replace(" ", "", regex=False),
        errors="coerce"
    )

# =========================================================
# FILTRO SIGMA
# =========================================================
def applica_filtro_sigma(serie, n_sigma=2.0):

    try:

        serie = converti_numerico(serie)

        media = serie.mean()
        std = serie.std()

        if pd.isna(std) or std == 0:
            return serie

        filtro = (
            (serie >= media - n_sigma * std) &
            (serie <= media + n_sigma * std)
        )

        return serie.where(filtro)

    except Exception as e:
        logger.error(f"Errore filtro sigma: {e}")
        return serie

# =========================================================
# CARICAMENTO EXCEL
# =========================================================
@st.cache_data
def carica_excel(uploaded_file):

    data = {}

    try:

        xl = pd.ExcelFile(
            uploaded_file,
            engine="openpyxl"
        )

        for sheet in xl.sheet_names:

            try:

                df = xl.parse(sheet, header=0)

                df = df.dropna(how="all")

                df.columns = [
                    str(c).strip()
                    for c in df.columns
                ]

                data[sheet] = df

            except Exception as e:
                logger.warning(f"Errore foglio {sheet}: {e}")

        return data

    except Exception as e:

        logger.error(f"Errore caricamento Excel: {e}")
        return {}

# =========================================================
# ESTRAZIONE COLONNE NUMERICHE
# =========================================================
def estrai_colonne_numeriche(df):

    colonne = []

    for c in df.columns[1:]:

        if "Unnamed" in str(c):
            continue

        serie_test = converti_numerico(df[c])

        if serie_test.notnull().sum() > 0:
            colonne.append(c)

    return colonne

# =========================================================
# CREAZIONE REPORT WORD
# =========================================================
def genera_report_word(
    punti,
    configurazione,
    metodo,
    n_sigma,
    metriche_globali,
    image_path
):

    doc = Document()

    # =====================================================
    # TITOLO
    # =====================================================
    doc.add_heading(
        'DIMOS - REPORT ANALISI TOPOGRAFICA',
        level=1
    )

    doc.add_paragraph(
        f"Metodo elaborazione: {metodo}"
    )

    if metodo == "Filtro Sigma (Gauss)":
        doc.add_paragraph(
            f"Valore Sigma: {n_sigma}"
        )

    doc.add_paragraph(
        f"Punti analizzati: {', '.join(punti)}"
    )

    # =====================================================
    # GRAFICO
    # =====================================================
    if os.path.exists(image_path):

        doc.add_heading(
            'Grafico Analisi',
            level=2
        )

        doc.add_picture(
            image_path,
            width=Inches(7)
        )

    # =====================================================
    # METRICHE
    # =====================================================
    doc.add_heading(
        'Metriche',
        level=2
    )

    table = doc.add_table(
        rows=1,
        cols=6
    )

    hdr_cells = table.rows[0].cells

    hdr_cells[0].text = "Punto"
    hdr_cells[1].text = "Parametro"
    hdr_cells[2].text = "MIN"
    hdr_cells[3].text = "MAX"
    hdr_cells[4].text = "RANGE"
    hdr_cells[5].text = "ULTIMO"

    for row in metriche_globali:

        cells = table.add_row().cells

        cells[0].text = row["punto"]
        cells[1].text = row["parametro"]
        cells[2].text = f"{row['min']:.3f}"
        cells[3].text = f"{row['max']:.3f}"
        cells[4].text = f"{row['range']:.3f}"
        cells[5].text = f"{row['ultimo']:.3f}"

    # =====================================================
    # SALVATAGGIO
    # =====================================================
    tmp_docx = tempfile.NamedTemporaryFile(
        delete=False,
        suffix=".docx"
    )

    doc.save(tmp_docx.name)

    return tmp_docx.name

# =========================================================
# APP PRINCIPALE
# =========================================================
def run_tps_monitoring():

    st.title("🛰️ DIMOS - Analisi Topografica Avanzata")

    st.markdown("""
    ### Funzioni Disponibili

    ✅ Analisi multilayer  
    ✅ Multi-sensore  
    ✅ Trendline automatiche  
    ✅ Filtro Sigma  
    ✅ Report Word automatico  
    ✅ Gestione Excel sporchi  
    ✅ Compatibile Streamlit Cloud
    """)

    # =====================================================
    # TABS
    # =====================================================
    tab1, tab2, tab3, tab4 = st.tabs([
        "📂 Caricamento",
        "⚙️ Configurazione",
        "📈 Analisi",
        "📄 Report"
    ])

    # =====================================================
    # TAB 1 - CARICAMENTO
    # =====================================================
    with tab1:

        uploaded_file = st.file_uploader(
            "Carica file Excel",
            type=["xlsx"]
        )

        if uploaded_file is not None:

            dfs = carica_excel(uploaded_file)

            st.success(
                f"Fogli trovati: {len(dfs)}"
            )

            st.session_state["dfs"] = dfs

            st.write(list(dfs.keys()))

    # =====================================================
    # VERIFICA DATI
    # =====================================================
    if "dfs" not in st.session_state:
        return

    dfs = st.session_state["dfs"]

    # =====================================================
    # TAB 2 - CONFIGURAZIONE
    # =====================================================
    with tab2:

        st.header("Configurazione Analisi")

        fogli = list(dfs.keys())

        punti = st.multiselect(
            "Seleziona punti/sensori",
            fogli
        )

        if len(punti) == 0:
            st.warning("Seleziona almeno un punto.")
            return

        metodo = st.radio(
            "Metodo elaborazione",
            [
                "Dati Completi",
                "Filtro Sigma (Gauss)"
            ]
        )

        n_sigma = 2.0

        if metodo == "Filtro Sigma (Gauss)":

            n_sigma = st.slider(
                "Valore Sigma",
                min_value=1.0,
                max_value=5.0,
                value=2.0,
                step=0.5
            )

        # =================================================
        # CONFIGURAZIONE MULTILAYER
        # =================================================
        configurazione = {}

        st.markdown("---")
        st.subheader("Selezione Parametri")

        for punto in punti:

            st.markdown(f"### 📍 {punto}")

            df = dfs[punto]

            colonne = estrai_colonne_numeriche(df)

            selezione = st.multiselect(
                f"Parametri disponibili - {punto}",
                colonne,
                default=colonne[:1],
                key=f"cfg_{punto}"
            )

            configurazione[punto] = selezione

        st.session_state["configurazione"] = configurazione
        st.session_state["punti"] = punti
        st.session_state["metodo"] = metodo
        st.session_state["n_sigma"] = n_sigma

    # =====================================================
    # TAB 3 - ANALISI
    # =====================================================
    with tab3:

        if "configurazione" not in st.session_state:
            return

        configurazione = st.session_state["configurazione"]
        metodo = st.session_state["metodo"]
        n_sigma = st.session_state["n_sigma"]

        fig = go.Figure()

        metriche_globali = []

        # =================================================
        # LOOP MULTILAYER
        # =================================================
        for punto in configurazione:

            df = dfs[punto].copy()

            if df.empty:
                continue

            col_data = df.columns[0]

            df[col_data] = pd.to_datetime(
                df[col_data],
                errors="coerce",
                dayfirst=True
            )

            df = df.dropna(
                subset=[col_data]
            ).sort_values(col_data)

            for parametro in configurazione[punto]:

                try:

                    d = df[[col_data, parametro]].copy()

                    d[parametro] = converti_numerico(
                        d[parametro]
                    )

                    d = d.dropna()

                    if metodo == "Filtro Sigma (Gauss)":

                        d[parametro] = applica_filtro_sigma(
                            d[parametro],
                            n_sigma
                        )

                    d = d.dropna()

                    if d.empty:
                        continue

                    # =====================================
                    # METRICHE
                    # =====================================
                    minimo = d[parametro].min()
                    massimo = d[parametro].max()
                    ultimo = d[parametro].iloc[-1]
                    delta = massimo - minimo

                    metriche_globali.append({
                        "punto": punto,
                        "parametro": parametro,
                        "min": minimo,
                        "max": massimo,
                        "range": delta,
                        "ultimo": ultimo
                    })

                    # =====================================
                    # GRAFICO DATI
                    # =====================================
                    fig.add_trace(
                        go.Scatter(
                            x=d[col_data],
                            y=d[parametro],
                            mode="lines+markers",
                            name=f"{punto} - {parametro}"
                        )
                    )

                    # =====================================
                    # RANGE VISIVO
                    # =====================================
                    fig.add_hrect(
                        y0=minimo,
                        y1=massimo,
                        opacity=0.05,
                        line_width=0
                    )

                    # =====================================
                    # TRENDLINE
                    # =====================================
                    if len(d) >= 5:

                        try:

                            x_calc = (
                                d[col_data] -
                                d[col_data].min()
                            ).dt.total_seconds() / 86400

                            coeff = np.polyfit(
                                x_calc,
                                d[parametro],
                                3
                            )

                            poly = np.poly1d(coeff)

                            fig.add_trace(
                                go.Scatter(
                                    x=d[col_data],
                                    y=poly(x_calc),
                                    mode="lines",
                                    name=f"Trend {punto}-{parametro}",
                                    line=dict(
                                        dash="dot",
                                        width=2
                                    )
                                )
                            )

                        except Exception as e:
                            logger.warning(
                                f"Trendline fallita: {e}"
                            )

                except Exception as e:

                    logger.error(
                        f"Errore {punto}-{parametro}: {e}"
                    )

        # =================================================
        # METRICHE
        # =================================================
        st.subheader("📊 Metriche")

        cols = st.columns(
            max(1, min(4, len(metriche_globali)))
        )

        for i, m in enumerate(metriche_globali):

            cols[i % len(cols)].metric(
                label=f"{m['punto']} | {m['parametro']}",
                value=f"{m['ultimo']:.3f}",
                delta=f"Range {m['range']:.3f}"
            )

        # =================================================
        # GRAFICO
        # =================================================
        fig.update_layout(
            template="plotly_white",
            height=750,
            hovermode="x unified",
            legend=dict(
                orientation="h",
                yanchor="bottom",
                y=1.02,
                xanchor="right",
                x=1
            ),
            xaxis=dict(
                title="Data",
                tickformat="%d/%m/%Y",
                tickangle=-45
            ),
            yaxis=dict(
                title="Valore"
            )
        )

        st.plotly_chart(
            fig,
            use_container_width=True
        )

        # =================================================
        # SALVATAGGIO FIGURA
        # =================================================
        img_tmp = tempfile.NamedTemporaryFile(
            delete=False,
            suffix=".png"
        )

        try:

            fig.write_image(
                img_tmp.name,
                width=1800,
                height=900
            )

            st.session_state["img_path"] = img_tmp.name

        except Exception as e:

            st.warning("""
            Installare KALEIDO per export immagini:

            pip install kaleido
            """)

            logger.error(e)

        st.session_state["metriche"] = metriche_globali

    # =====================================================
    # TAB 4 - REPORT
    # =====================================================
    with tab4:

        st.header("📄 Esportazione Report")

        if st.button("Genera Report Word"):

            try:

                report_path = genera_report_word(
                    punti=st.session_state["punti"],
                    configurazione=st.session_state["configurazione"],
                    metodo=st.session_state["metodo"],
                    n_sigma=st.session_state["n_sigma"],
                    metriche_globali=st.session_state["metriche"],
                    image_path=st.session_state["img_path"]
                )

                with open(report_path, "rb") as file:

                    st.download_button(
                        label="⬇️ Scarica Report Word",
                        data=file,
                        file_name="DIMOS_REPORT.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                st.success("Report generato.")

            except Exception as e:

                st.error(f"Errore report: {e}")
                logger.exception(e)

# =========================================================
# AVVIO
# =========================================================
if __name__ == "__main__":
    run_tps_monitoring()
