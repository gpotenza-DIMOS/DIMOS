import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from io import BytesIO
from datetime import datetime
import os

# --- GESTIONE LIBRERIA DOCX (REPORT WORD) ---
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# --- LOGICA FILTRI (GAUSS E ZERI) ---
def applica_filtri_completi(serie, n_sigma, rimuovi_zeri):
    originale = serie.copy()
    diag = {"zeri": 0, "gauss": 0}
    
    # 1. Rimozione Zeri Puri
    if rimuovi_zeri:
        diag["zeri"] = (originale == 0).sum()
        originale = originale.replace(0, np.nan)
    
    # 2. Filtro Gauss (Outliers)
    validi = originale.dropna()
    if not validi.empty and n_sigma > 0:
        mean, std = validi.mean(), validi.std()
        if std > 0:
            lower, upper = mean - n_sigma * std, mean + n_sigma * std
            outliers = (originale < lower) | (originale > upper)
            diag["gauss"] = outliers.sum()
            originale[outliers] = np.nan
            
    return originale, diag

# --- FUNZIONE CARICAMENTO EXCEL (ROBUSTA) ---
@st.cache_resource
def get_data_from_excel(file_content):
    try:
        xls = pd.ExcelFile(file_content)
        # Legge foglio NAME per la struttura sensori
        df_header = pd.read_excel(xls, sheet_name='NAME', header=None, nrows=3)
        
        # Legge il primo foglio dati (quello dopo NAME)
        data_sheet = [s for s in xls.sheet_names if s != 'NAME'][0]
        df_values = pd.read_excel(xls, sheet_name=data_sheet)
        
        # Pulizia nomi colonne (rimuove spazi bianchi)
        df_values.columns = [str(c).strip() for c in df_values.columns]
        
        # RICERCA INTELLIGENTE COLONNA TEMPORALE (Risolve KeyError)
        col_tempo = None
        percorsi_data = ['data', 'time', 'ora', 'date', 'timestamp']
        for c in df_values.columns:
            if any(p in c.lower() for p in percorsi_data):
                col_tempo = c
                break
        
        if col_tempo:
            df_values[col_tempo] = pd.to_datetime(df_values[col_tempo], errors='coerce')
            df_values = df_values.dropna(subset=[col_tempo]).sort_values(by=col_tempo)
            return df_header, df_values, col_tempo
        else:
            return None, None, None
    except Exception as e:
        return None, None, None

# --- GENERAZIONE REPORT WORD ---
def genera_report_word(fig, stats, tipo_sel, d_range, sigma, zeri_on):
    doc = Document()
    
    # Logo (se esiste)
    logo_path = "logo_dimos.jpg"
    if os.path.exists(logo_path):
        try: doc.add_picture(logo_path, width=Inches(1.5))
        except: pass
    
    doc.add_heading('REPORT MONITORAGGIO DIMOS', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run(f"Data Report: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n").bold = True
    p.add_run(f"Periodo Analizzato: {d_range[0]} - {d_range[1]}\n")
    p.add_run(f"Filtri: Gauss ({sigma} σ), Zeri ({'Attivo' if zeri_on else 'No'})")

    # Immagine grafico (Kaleido richiesto nei requirements)
    try:
        img_bytes = fig.to_image(format="png", width=1000, height=500, scale=2)
        doc.add_picture(BytesIO(img_bytes), width=Inches(6))
    except Exception as e:
        doc.add_paragraph(f"\n[Grafico non disponibile: {e}]\n")

    doc.add_heading('Diagnostica Dati', level=1)
    table = doc.add_table(rows=1, cols=3); table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'Sensore/Asse', 'Zeri Rimossi', 'Outliers Gauss'
    for s_name, s_diag in stats.items():
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = s_name, str(s_diag['zeri']), str(s_diag['gauss'])
    
    target = BytesIO()
    doc.save(target)
    return target.getvalue()

# --- FUNZIONE PRINCIPALE INTERFACCIA ---
def run_plotter():
    st.header("📉 PLOTTER PRO - Analisi e Strutture")
    
    with st.sidebar:
        st.header("⚙️ Impostazioni Asse X")
        tipo_asse_x = st.radio("Modalità:", ["Sequenziale (Stile Excel)", "Temporale"])
        passo_date = st.number_input("Etichette ogni N date:", value=400, min_value=1)
        st.divider()
        st.header("🔍 Filtri")
        rimuovi_zeri = st.toggle("Rimuovi Zeri Puri", value=True)
        usa_gauss = st.checkbox("Filtro Gauss (Spike)", value=True)
        sigma = st.slider("Sensibilità Sigma", 0.5, 5.0, 2.0) if usa_gauss else 0
        
    file_input = st.sidebar.file_uploader("Carica Excel", type=['xlsx', 'xlsm'], key="plt_v_final_top")
    
    if not file_input:
        st.info("👋 Benvenuto! Carica un file Excel dal menu a sinistra per iniziare l'analisi.")
        return

    # Esecuzione solo se il file è presente
    df_header, df_values, col_tempo = get_data_from_excel(file_input)
    
    if df_values is None or col_tempo is None:
        st.error("⚠️ Impossibile trovare la colonna temporale nel file. Verifica che si chiami 'Data' o 'Data e Ora'.")
        return

    # 1. Mappatura Sensori
    mappa = {}
    for col in range(1, len(df_header.columns)):
        id_t, nome, logger = str(df_header.iloc[2, col]), str(df_header.iloc[1, col]), str(df_header.iloc[0, col])
        # Identificazione tipo grandezza
        if "[°]" in id_t: t = "Inclinazione (°)"
        elif "°C" in id_t.upper(): t = "Temperatura (°C)"
        elif "MM" in id_t.upper(): t = "Spostamento (mm)"
        elif "[V]" in id_t.upper(): t = "Batteria (V)"
        else: t = "Altro"
        
        label_sensore = f"{nome} ({logger})"
        if label_sensore not in mappa: mappa[label_sensore] = {"tipo": t, "canali": {}}
        
        # Estrazione asse (X, Y, Z) dall'ID tecnico
        suffix = id_t.split()[-2] if len(id_t.split()) > 2 else id_t
        mappa[label_sensore]["canali"][id_t] = suffix

    # 2. Selezione Dinamica
    grandezze_disp = sorted(list(set([v['tipo'] for v in mappa.values()])))
    tipo_sel = st.selectbox("Seleziona Grandezza Fisica:", grandezze_disp)
    
    sensori_disp = [k for k, v in mappa.items() if v['tipo'] == tipo_sel]
    sensori_scelti = st.multiselect("Seleziona Sensori da visualizzare:", sensori_disp)

    final_targets = {}
    if sensori_scelti:
        st.write("### 🎯 Selezione Assi e Canali")
        cols = st.columns(len(sensori_scelti))
        for i, s in enumerate(sensori_scelti):
            with cols[i]:
                canali = mappa[s]["canali"]
                scelte = st.multiselect(f"{s}:", list(canali.keys()), default=list(canali.keys())[:1], format_func=lambda x: canali[x])
                for sc in scelte:
                    final_targets[f"{s} - {canali[sc]}"] = sc

    # 3. Grafico e Diagnostica
    if final_targets:
        st.divider()
        c1, c2 = st.columns(2)
        with c1: start_d = st.date_input("Data Inizio:", df_values[col_tempo].min().date())
        with c2: end_d = st.date_input("Data Fine:", df_values[col_tempo].max().date())
        
        if start_d <= end_d:
            # Filtro temporale
            mask = (df_values[col_tempo].dt.date >= start_d) & (df_values[col_tempo].dt.date <= end_d)
            df_p = df_values.loc[mask].copy().reset_index(drop=True)
            
            fig = go.Figure()
            stats_final = {}
            
            for label, cid in final_targets.items():
                if cid in df_p.columns:
                    y_clean, diag = applica_filtri_completi(df_p[cid], sigma, rimuovi_zeri)
                    stats_final[label] = diag
                    
                    # Sincronizzazione per asse X (evita errori di lunghezza)
                    temp_df = pd.DataFrame({'idx': df_p.index, 'time': df_p[col_tempo], 'y': y_clean}).dropna(subset=['y'])
                    x_axis = temp_df['idx'] if "Sequenziale" in tipo_asse_x else temp_df['time']
                    
                    fig.add_trace(go.Scatter(x=x_axis, y=temp_df['y'], name=label, mode='lines+markers', connectgaps=True))

            # Formattazione Asse X stile Excel (Etichetta ogni N passi)
            if "Sequenziale" in tipo_asse_x:
                ticks_txt = [row[col_tempo].strftime('%d/%m/%y %H:%M') if i % passo_date == 0 else "" for i, row in df_p.iterrows()]
                fig.update_layout(xaxis=dict(type='category', tickmode='array', tickvals=list(df_p.index), ticktext=ticks_txt, tickangle=-45))
            else:
                fig.update_layout(xaxis=dict(type='date', tickangle=-45))

            fig.update_layout(template="plotly_white", height=650, hovermode="x unified", legend=dict(orientation="h", yanchor="bottom", y=1.02))
            st.plotly_chart(fig, use_container_width=True)
            
            # 4. Diagnostica e Esportazione
            cd1, cd2 = st.columns([2, 1])
            with cd1:
                with st.expander("📊 Diagnostica Filtri (Valori scartati)"):
                    st.table(pd.DataFrame(stats_final).T.rename(columns={"zeri": "Zeri", "gauss": "Gauss"}))
            with cd2:
                if st.button("📝 Genera Report Word"):
                    if DOCX_AVAILABLE:
                        with st.spinner("Preparazione report..."):
                            doc_b = genera_report_word(fig, stats_final, tipo_sel, [start_d, end_d], sigma, rimuovi_zeri)
                            st.download_button("📥 Scarica Report", doc_b, f"Report_{tipo_sel}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    elif not final_targets and file_input:
        st.warning("Seleziona almeno un sensore e un asse per visualizzare i dati.")

if __name__ == "__main__":
    run_plotter()
