import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import warnings
import logging
import tempfile
import os

from docx import Document
from docx.shared import Inches

# =========================================================
# CONFIGURAZIONE E UTILITY
# =========================================================
warnings.filterwarnings("ignore")
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger("DIMOS")

def converti_numerico(serie):
    return pd.to_numeric(
        serie.astype(str).str.replace(",", ".", regex=False).str.replace(" ", "", regex=False),
        errors="coerce"
    )

def applica_filtro_sigma(serie, n_sigma=2.0):
    try:
        serie = converti_numerico(serie)
        media = serie.mean()
        std = serie.std()
        if pd.isna(std) or std == 0:
            return serie
        filtro = ((serie >= media - n_sigma * std) & (serie <= media + n_sigma * std))
        return serie.where(filtro)
    except Exception as e:
        logger.error(f"Errore filtro sigma: {e}")
        return serie

@st.cache_data
def carica_excel(uploaded_file):
    data = {}
    try:
        xl = pd.ExcelFile(uploaded_file, engine="openpyxl")
        for sheet in xl.sheet_names:
            df = xl.parse(sheet, header=0)
            df = df.dropna(how="all")
            df.columns = [str(c).strip() for c in df.columns]
            data[sheet] = df
        return data
    except Exception as e:
        logger.error(f"Errore caricamento Excel: {e}")
        return {}

def estrai_colonne_numeriche(df):
    colonne = []
    for c in df.columns[1:]:
        if "Unnamed" in str(c): continue
        serie_test = converti_numerico(df[c])
        if serie_test.notnull().sum() > 0:
            colonne.append(c)
    return colonne

def ottieni_default_params(colonne_disponibili):
    """Ritorna DeltaE e DeltaN se presenti, altrimenti la prima colonna."""
    defaults = [c for c in colonne_disponibili if c.lower() in ["deltae", "deltan", "delta e", "delta n"]]
    if not defaults and colonne_disponibili:
        defaults = [colonne_disponibili[0]]
    return defaults

# =========================================================
# CREAZIONE REPORT WORD
# =========================================================
def genera_report_word(metodo, n_sigma, metriche_globali, image_path):
    doc = Document()
    doc.add_heading('DIMOS - REPORT ANALISI TOPOGRAFICA', level=1)
    
    doc.add_paragraph(f"Metodo elaborazione: {metodo}")
    if metodo == "Filtro Sigma (Gauss)":
        doc.add_paragraph(f"Valore Sigma: {n_sigma}")

    # Sezione Grafico
    if image_path and os.path.exists(image_path):
        doc.add_heading('Visualizzazione Dati (Report)', level=2)
        doc.add_picture(image_path, width=Inches(6.5))

    # Sezione Metriche
    doc.add_heading('Tabella Metriche', level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(["Punto", "Parametro", "MIN", "MAX", "RANGE", "ULTIMO"]):
        hdr_cells[i].text = text

    for row in metriche_globali:
        cells = table.add_row().cells
        cells[0].text = str(row["punto"])
        cells[1].text = str(row["parametro"])
        cells[2].text = f"{row['min']:.3f}"
        cells[3].text = f"{row['max']:.3f}"
        cells[4].text = f"{row['range']:.3f}"
        cells[5].text = f"{row['ultimo']:.3f}"

    tmp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp_docx.name)
    return tmp_docx.name

# =========================================================
# APP PRINCIPALE
# =========================================================
def run_tps_monitoring():
    st.set_page_config(page_title="DIMOS - Analisi Topografica", layout="wide")
    st.title("🛰️ DIMOS - Analisi Topografica Avanzata")
    
    if "uploaded_file" not in st.session_state:
        uploaded_file = st.file_uploader("📂 Carica file Excel (.xlsx)", type=["xlsx"])
    else:
        uploaded_file = st.session_state.uploaded_file

    if uploaded_file is not None:
        dfs = carica_excel(uploaded_file)
        fogli = list(dfs.keys())
        
        # Estrazione di tutti i parametri possibili da tutti i fogli per i default
        tutte_colonne = []
        for f in fogli:
            tutte_colonne.extend(estrai_colonne_numeriche(dfs[f]))
        tutte_colonne = sorted(list(set(tutte_colonne)))

        # --- 2. CONFIGURAZIONE VISUALIZZAZIONE A VIDEO ---
        st.subheader("📺 Configurazione Visualizzazione (Grafico a Video)")
        with st.container(border=True):
            cv1, cv2 = st.columns([2, 1])
            with cv1:
                punti_video = st.multiselect(
                    "Seleziona Sensori per la visualizzazione", 
                    fogli, 
                    key="punti_video"
                )
            with cv2:
                metodo = st.radio("Metodo elaborazione", ["Dati Completi", "Filtro Sigma (Gauss)"], horizontal=True)
                n_sigma = 2.0
                if metodo == "Filtro Sigma (Gauss)":
                    n_sigma = st.slider("Valore Sigma", 1.0, 5.0, 2.0, 0.5)

            if punti_video:
                parametri_video = st.multiselect(
                    "Seleziona Parametri da visualizzare (uguali per tutti i sensori)",
                    tutte_colonne,
                    default=ottieni_default_params(tutte_colonne),
                    key="params_video"
                )

        # --- 3. ELABORAZIONE GRAFICO VIDEO ---
        if punti_video and parametri_video:
            fig_video = go.Figure()
            metriche_video = []

            for punto in punti_video:
                df = dfs[punto].copy()
                col_data = df.columns[0]
                df[col_data] = pd.to_datetime(df[col_data], errors="coerce", dayfirst=True)
                df = df.dropna(subset=[col_data]).sort_values(col_data)

                for parametro in parametri_video:
                    if parametro not in df.columns: continue
                    
                    d = df[[col_data, parametro]].copy()
                    d[parametro] = converti_numerico(d[parametro])
                    d = d.dropna()

                    if metodo == "Filtro Sigma (Gauss)":
                        d[parametro] = applica_filtro_sigma(d[parametro], n_sigma)
                        d = d.dropna()

                    if d.empty: continue

                    minimo, massimo = d[parametro].min(), d[parametro].max()
                    ultimo = d[parametro].iloc[-1]
                    metriche_video.append({
                        "punto": punto, "parametro": parametro,
                        "min": minimo, "max": massimo, "range": massimo - minimo, "ultimo": ultimo
                    })

                    fig_video.add_trace(go.Scatter(
                        x=d[col_data], y=d[parametro],
                        mode="lines+markers", name=f"{punto}: {parametro}"
                    ))

            # Render Grafico Video
            st.plotly_chart(fig_video, use_container_width=True)
            
            # Espandibile per metriche a video
            with st.expander("📊 Visualizza Metriche Real-time"):
                m_cols = st.columns(4)
                for idx, m in enumerate(metriche_video):
                    m_cols[idx % 4].metric(
                        label=f"{m['punto']} | {m['parametro']}",
                        value=f"{m['ultimo']:.3f}",
                        delta=f"R: {m['range']:.3f}"
                    )
        else:
            st.info("Seleziona sensori e parametri per visualizzare il grafico.")

        # --- 4. CONFIGURAZIONE REPORT WORD (SEPARATA) ---
        st.divider()
        st.subheader("📄 Configurazione Report Word")
        
        with st.container(border=True):
            cw1, cw2 = st.columns([2, 1])
            with cw1:
                punti_word = st.multiselect(
                    "Seleziona Sensori per il Report Word", 
                    ["TUTTI"] + fogli,
                    default=punti_video if punti_video else None,
                    key="punti_word"
                )
            with cw2:
                parametri_word = st.multiselect(
                    "Seleziona Parametri per il Report Word",
                    tutte_colonne,
                    default=ottieni_default_params(tutte_colonne),
                    key="params_word"
                )

        if st.button("🚀 Genera e Scarica Report Word"):
            if not punti_word or not parametri_word:
                st.error("Seleziona almeno un sensore e un parametro per il report.")
            else:
                with st.spinner("Generazione report indipendente in corso..."):
                    # Se selezionato TUTTI, espandiamo la lista
                    punti_da_stampare = fogli if "TUTTI" in punti_word else punti_word
                    
                    fig_word = go.Figure()
                    metriche_word = []

                    for punto in punti_da_stampare:
                        if punto not in dfs: continue
                        df = dfs[punto].copy()
                        col_data = df.columns[0]
                        df[col_data] = pd.to_datetime(df[col_data], errors="coerce", dayfirst=True)
                        df = df.dropna(subset=[col_data]).sort_values(col_data)

                        for parametro in parametri_word:
                            if parametro not in df.columns: continue
                            d = df[[col_data, parametro]].copy()
                            d[parametro] = converti_numerico(d[parametro])
                            d = d.dropna()
                            if metodo == "Filtro Sigma (Gauss)":
                                d[parametro] = applica_filtro_sigma(d[parametro], n_sigma)
                                d = d.dropna()
                            
                            if d.empty: continue
                            
                            metriche_word.append({
                                "punto": punto, "parametro": parametro,
                                "min": d[parametro].min(), "max": d[parametro].max(), 
                                "range": d[parametro].max() - d[parametro].min(), "ultimo": d[parametro].iloc[-1]
                            })
                            fig_word.add_trace(go.Scatter(x=d[col_data], y=d[parametro], mode="lines+markers", name=f"{punto}: {parametro}"))

                    # Salvataggio immagine e generazione documento
                    img_tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                    try:
                        fig_word.write_image(img_tmp.name, width=1200, height=600)
                        report_path = genera_report_word(metodo, n_sigma, metriche_word, img_tmp.name)
                        
                        with open(report_path, "rb") as f:
                            st.download_button(
                                label="⬇️ Clicca qui per scaricare il file Word",
                                data=f,
                                file_name="Report_Topografico_DIMOS.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    except Exception as e:
                        st.error(f"Errore: {e}. Verifica l'installazione di 'kaleido'.")

    else:
        st.info("Benvenuto in DIMOS. Carica un file Excel per iniziare l'analisi.")

if __name__ == "__main__":
    run_tps_monitoring()
